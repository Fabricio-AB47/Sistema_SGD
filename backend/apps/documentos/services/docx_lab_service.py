from __future__ import annotations

import json
import shutil
import uuid
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from django.conf import settings
from django.utils import timezone
from docx import Document


DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
DOCX_RESULT_ROOT = Path(getattr(settings, "MEDIA_ROOT", Path(settings.BASE_DIR) / "media")) / "docx_lab"
DEFAULT_DOCX_NAME = "documento.docx"


def _safe_name(name: str) -> str:
    candidate = Path(name or DEFAULT_DOCX_NAME).name.strip()
    if not candidate.lower().endswith(".docx"):
        candidate = f"{candidate}.docx"
    return candidate or DEFAULT_DOCX_NAME


def _write_uploaded_file(uploaded_file, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with destination.open("wb") as target:
        for chunk in uploaded_file.chunks():
            target.write(chunk)


def _read_comment_text(element) -> str:
    pieces = [node.text for node in element.findall(".//w:t", DOCX_NS) if node.text]
    return "".join(pieces).strip()


def _extract_comments_map(docx_path: Path) -> dict[str, dict[str, str]]:
    comments: dict[str, dict[str, str]] = {}
    with ZipFile(docx_path) as archive:
        if "word/comments.xml" not in archive.namelist():
            return comments

        root = ET.fromstring(archive.read("word/comments.xml"))
        for comment in root.findall("w:comment", DOCX_NS):
            comment_id = comment.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
            if not comment_id:
                continue
            comments[comment_id] = {
                "id": comment_id,
                "author": comment.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", ""),
                "initials": comment.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}initials", ""),
                "date": comment.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date", ""),
                "text": _read_comment_text(comment),
            }
    return comments


def _extract_paragraph_comment_ids(paragraph) -> list[str]:
    try:
        root = ET.fromstring(paragraph._p.xml)
    except ET.ParseError:
        return []

    comment_ids: list[str] = []
    for node in root.findall(".//w:commentReference", DOCX_NS):
        comment_id = node.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
        if comment_id and comment_id not in comment_ids:
            comment_ids.append(comment_id)
    return comment_ids


def _extract_tables(document: Document) -> list[dict[str, object]]:
    tables: list[dict[str, object]] = []
    for index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(
            {
                "index": index,
                "rows": rows,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
            }
        )
    return tables


def _extract_document_metadata(source_path: Path, uploaded_file, document: Document) -> dict[str, object]:
    stat = source_path.stat()
    return {
        "file_name": source_path.name,
        "file_size": stat.st_size,
        "mime_type": getattr(uploaded_file, "content_type", None),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "comment_count": 0,
        "created_at": timezone.now().isoformat(),
    }


def _copy_paragraph_format(source_paragraph, target_paragraph) -> None:
    if source_paragraph.style is not None:
        try:
            target_paragraph.style = source_paragraph.style.name
        except Exception:
            pass
    target_paragraph.alignment = source_paragraph.alignment


def _rebuild_docx_copy(source_path: Path, target_path: Path) -> None:
    source_document = Document(str(source_path))
    target_document = Document()

    if source_document.paragraphs:
        first_source_paragraph = source_document.paragraphs[0]
        first_target_paragraph = target_document.add_paragraph()
        _copy_paragraph_format(first_source_paragraph, first_target_paragraph)
        first_target_paragraph.text = first_source_paragraph.text

        for source_paragraph in source_document.paragraphs[1:]:
            target_paragraph = target_document.add_paragraph()
            _copy_paragraph_format(source_paragraph, target_paragraph)
            target_paragraph.text = source_paragraph.text

    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_document.save(str(target_path))


def build_docx_lab_result(*, uploaded_file) -> dict:
    DOCX_RESULT_ROOT.mkdir(parents=True, exist_ok=True)
    result_token = uuid.uuid4().hex
    result_dir = DOCX_RESULT_ROOT / result_token
    result_dir.mkdir(parents=True, exist_ok=True)

    file_name = _safe_name(getattr(uploaded_file, "name", "documento.docx"))
    source_path = result_dir / f"source_{file_name}"
    reconstructed_path = result_dir / f"reconstruido_{file_name}"

    _write_uploaded_file(uploaded_file, source_path)
    source_document = Document(str(source_path))
    comments_map = _extract_comments_map(source_path)

    paragraphs = []
    comment_ids_used: list[str] = []
    for index, paragraph in enumerate(source_document.paragraphs, start=1):
        paragraph_comment_ids = _extract_paragraph_comment_ids(paragraph)
        comment_ids_used.extend([comment_id for comment_id in paragraph_comment_ids if comment_id not in comment_ids_used])
        paragraphs.append(
            {
                "index": index,
                "style": getattr(getattr(paragraph, "style", None), "name", "Normal") or "Normal",
                "text": paragraph.text,
                "has_comment": bool(paragraph_comment_ids),
                "comment_ids": paragraph_comment_ids,
                "comments": [comments_map[comment_id] for comment_id in paragraph_comment_ids if comment_id in comments_map],
            }
        )

    tables = _extract_tables(source_document)
    metadata = _extract_document_metadata(source_path, uploaded_file, source_document)
    metadata["comment_count"] = len(comments_map)
    metadata["paragraph_count"] = len(paragraphs)
    metadata["table_count"] = len(tables)
    metadata["has_comments"] = bool(comments_map)
    metadata["has_tables"] = bool(tables)
    metadata["has_content"] = bool(paragraphs or tables)

    _rebuild_docx_copy(source_path, reconstructed_path)

    analysis = {
        "generated_at": timezone.now().isoformat(),
        "source_name": file_name,
        "source_path": source_path.as_posix(),
        "reconstructed_path": reconstructed_path.as_posix(),
        "paragraph_count": len(paragraphs),
        "comment_count": len(comments_map),
        "comment_ids_used": comment_ids_used,
        "comments": list(comments_map.values()),
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": metadata,
    }

    analysis_path = result_dir / "analysis.json"
    analysis_path.write_text(json.dumps(analysis, ensure_ascii=False, indent=2), encoding="utf-8")

    return {
        "result_token": result_token,
        "result_dir": result_dir,
        "analysis_path": analysis_path,
        "analysis": analysis,
        "source_path": source_path,
        "reconstructed_path": reconstructed_path,
        "source_name": file_name,
    }


def load_docx_lab_result(result_token: str | None) -> dict | None:
    if not result_token:
        return None

    analysis_path = DOCX_RESULT_ROOT / result_token / "analysis.json"
    if not analysis_path.exists():
        return None
    return json.loads(analysis_path.read_text(encoding="utf-8"))


def cleanup_docx_lab_result(result_token: str | None) -> None:
    if not result_token:
        return

    result_dir = DOCX_RESULT_ROOT / result_token
    if result_dir.exists():
        shutil.rmtree(result_dir, ignore_errors=True)